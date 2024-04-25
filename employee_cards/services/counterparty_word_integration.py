import calendar
import datetime
import io
import json
from datetime import date
from string import Template
from typing import Any

from django.db.models import Sum

from docx import Document
from docx.document import Document as DocObject
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from employee_cards.models import Counterparty

from global_utils import num2text


class CounterpartiesWordIntegrationService:
    group_k_convert: dict[str, str] = {
        "Поставщики": "Поставщиков",
        "Государственные органы": "Государственных органов",
        "Покупатели": "Покупателей",
        "Проблемные поставщики": "Проблемных поставщиков",
    }

    type_k_convert: dict[str, str] = {
        "Физическое лицо": "Физическим лицом",
        "Юридическое лицо": "Юридическим лицом",
    }

    current_year: int = date.today().year

    def __init__(self, font: str = "Arial", font_size: int = 14) -> None:
        self.__is_generated = False
        self.__init_counterparties()
        self.__init_document(font, font_size)
        self.__read_templates()
        self.__abc_analysis()

    def __init_counterparties(self) -> None:
        self.__counterparties = (
            Counterparty.objects.all()
            .order_by("-total")
            .prefetch_related("invoices", "changing_logs")
        )

    def __init_document(self, font: str, font_size: int) -> None:
        self.document: DocObject = Document()
        self.document.styles["Normal"].font.name = font
        self.document.styles["Normal"].font.size = Pt(font_size)

    def __read_templates(self) -> None:
        with open("data/counterparties/reports_templates.json", "r") as f:
            data = json.load(f)
            self.__templates: dict[str, Template] = {}
            for key, template in data.items():
                self.__templates[key] = Template(template)

    def __abc_analysis(self) -> None:
        self.__sum_totals = Counterparty.objects.aggregate(Sum("total"))[
            "total__sum"
        ]
        self.__abc_data: dict[str, str] = {}

        total_percent = 0
        for counterparty in self.__counterparties:
            total = counterparty.total
            if total is None:
                total = 0
            curr_percent = (total * 100) / self.__sum_totals
            total_percent += curr_percent

            if total_percent < 80:
                self.__abc_data[counterparty.inn] = "А"
            elif total_percent < 95:
                self.__abc_data[counterparty.inn] = "Б"
            else:
                self.__abc_data[counterparty.inn] = "С"

    def generate_report(self) -> None:
        self.__add_header_to_document()

        for i, counterparty in enumerate(self.__counterparties):
            paragraph = self.__templates["counterparty_start"].substitute(
                self.__get_counterparty_info(counterparty, i)
            )

            paragraph += self.__templates[
                "counterparty_totals_in_months"
            ].substitute(self.__get_month_totals(counterparty))

            paragraph += self.__get_price_logs(counterparty)

            self.document.add_paragraph(
                paragraph
            ).alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE

        self.__is_generated = True

    def save_report(self) -> io.BytesIO:
        if not self.__is_generated:
            self.generate_report()
        file_stream = io.BytesIO()
        self.document.save(file_stream)
        file_stream.seek(0)
        return file_stream

    @staticmethod
    def __convert_date_to_string(init_date: datetime.datetime | None) -> str:
        if init_date:
            return init_date.strftime("%d.%m.%Y")
        return "<ДАТА НЕ ОПРЕДЕЛЕНА>"

    def __add_header_to_document(self) -> None:
        self.document.add_paragraph(
            self.__templates["header"].substitute()
        ).alignment = WD_ALIGN_PARAGRAPH.DISTRIBUTE

    def __get_counterparty_info(
        self, counterparty: Counterparty, index: int
    ) -> dict[str, Any]:
        group_k = counterparty.group_k
        type_k = counterparty.type_k
        total = counterparty.total
        stdev = counterparty.standard_deviation
        load_by_month = counterparty.load_by_month
        if not group_k:
            group_k = ""
        if not type_k:
            type_k = ""
        if not total:
            total = 0
        if not stdev:
            stdev = 0
        if not load_by_month:
            load_by_month = 0

        return {
            "number": num2text(index + 1),
            "name": counterparty.title,
            "group_abc": self.__abc_data.get(counterparty.inn, "С"),
            "inn": counterparty.inn,
            "group_k": self.group_k_convert.get(
                group_k, "<ГРУППА НЕ ОПРЕДЕЛЕНА>"
            ),
            "type_k": self.type_k_convert.get(
                type_k, "<ГРУППА НЕ ОПРЕДЕЛЕНА>"
            ),
            "init_date": self.__convert_date_to_string(counterparty.date_time),
            "total": round(total, 2),
            "stdev": round(stdev, 2),
            "load_by_month": round(load_by_month, 2),
        }

    def __get_month_totals(
        self, counterparty: Counterparty
    ) -> dict[str, float]:
        month_totals: dict[int, float] = {i: 0 for i in range(1, 13)}
        for invoice in counterparty.invoices.filter(
            date__year=self.current_year
        ):
            month = invoice.date.month
            month_totals[month] = month_totals.get(month, 0) + invoice.total
        return {
            calendar.month_name[key].lower(): round(value, 2)
            for key, value in month_totals.items()
        }

    def __get_price_logs(self, counterparty: Counterparty) -> str:
        paragraph = ""

        logs = counterparty.changing_logs.all()
        if logs.exists():
            paragraph += self.__templates["invoices_list_start"].substitute(
                {
                    "name": counterparty.title,
                }
            )

            for price_log in logs:
                nomenclature = price_log.nomenclature
                amount = round(price_log.amount)
                min_price = round(price_log.min_price, 2)
                max_price = round(price_log.max_price, 2)
                price_diff = round(
                    price_log.max_price - price_log.min_price, 2
                )

                placeholders_dict: dict[str, Any] = {
                    "nomenclature": nomenclature,
                    "amount": amount,
                    "price": max_price,
                }

                if abs(price_log.min_price - price_log.max_price) <= 11:
                    paragraph += self.__templates[
                        "invoices_list_no_change"
                    ].substitute(placeholders_dict)
                else:
                    paragraph += self.__templates[
                        "invoices_list_change"
                    ].substitute(
                        {
                            **placeholders_dict,
                            "price_diff": price_diff,
                            "min_price": min_price,
                        }
                    )

        return paragraph
